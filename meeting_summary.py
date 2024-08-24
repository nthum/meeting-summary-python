from openai import OpenAI
import streamlit as st
from docx import Document
import os

client = OpenAI(
    # defaults to os.environ.get("OPENAI_API_KEY")
    api_key = "OPENAI_API_KEY" ,
)

# Function to transcribe audio file using OpenAI Speech-to-Text API
def transcribe_audio(audio_file):
    transcription = client.audio.transcriptions.create(model="whisper-1", file=audio_file)
    return transcription.text

def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

def abstract_summary_extraction(transcription):
    completion = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ],
        max_tokens=100
    )
    return completion.choices[0].message.content

def key_points_extraction(transcription):
    completion = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return completion.choices[0].message.content

def action_item_extraction(transcription):
    completion = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return completion.choices[0].message.content

def sentiment_analysis(transcription):
    completion = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return completion.choices[0].message.content

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

    # Streamlit UI
def main():

    combined_transcript = ""
    
    st.title("AI-Powered Meeting Summarization Tool")

    uploaded_files = st.file_uploader("Upload an audio/video files (Limit 25MB per file) for transcription:", type=None, accept_multiple_files=True)

    if uploaded_files:
        for uploaded_file in uploaded_files:
            st.audio(uploaded_file, format='audio/mp3')
        
        if st.button("Transcribe"):
            try: 
                st.header("Processing uploaded files...")
                summaries = []
                for uploaded_file in uploaded_files:
                    st.header(f"Processing {uploaded_file.name}...")
                                    
                    # Transcribe the audio file
                    transcript = transcribe_audio(uploaded_file)
                    st.text(f"Transcript for {uploaded_file.name}:")
                    st.write(transcript)
                    
                    # Generate summary
                    summary = abstract_summary_extraction(transcript)
                    # st.text(f"Summary for {uploaded_file.name}:")
                    # st.write(summary)
                    
                    # Store summary for final overview
                    summaries.append({
                        'filename': uploaded_file.name,
                        'summary': summary
                    })

                    combined_transcript += transcript

                # generate an aggregate summary
                aggregate_summary = abstract_summary_extraction(combined_transcript)
                st.header("Summary:")
                st.write(aggregate_summary)

                key_points = key_points_extraction(combined_transcript)
                st.header("Key Points:")
                st.write(key_points)

                action_items = action_item_extraction(combined_transcript)
                st.header("Action Items:")
                st.write(action_items)

                sentiment = sentiment_analysis(combined_transcript)
                st.header("Sentiment Analysis:")
                st.write(sentiment)

            except Exception as e:
                st.error(f"Error: {e}")

if __name__ == "__main__":
    main()